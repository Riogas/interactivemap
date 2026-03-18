#!/usr/bin/env python3
"""
Genera documentación profesional de TrackMóvil en formato .docx
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── Estilos globales ──────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
    h.font.name = 'Calibri'

# ── Funciones auxiliares ──────────────────────────────────────
def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    # Rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    return table

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(f' {text}')
    else:
        p.add_run(text)

# ══════════════════════════════════════════════════════════════
# PORTADA
# ══════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('TrackMóvil')
run.font.size = Pt(42)
run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Sistema de Rastreo Vehicular en Tiempo Real')
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('Documentación Técnica y Funcional')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

doc.add_paragraph()
ver = doc.add_paragraph()
ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = ver.add_run('Versión 1.0  •  Julio 2025')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

client = doc.add_paragraph()
client.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = client.add_run('Riogas S.A.')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
run.bold = True

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# ÍNDICE
# ══════════════════════════════════════════════════════════════
doc.add_heading('Índice', level=1)
toc_items = [
    '1. Introducción',
    '2. Arquitectura del Sistema',
    '3. Autenticación y Control de Acceso',
    '4. Mapa Interactivo',
    '5. Gestión de Flota Vehicular',
    '6. Gestión de Pedidos',
    '7. Gestión de Servicios',
    '8. Sistema de Zonas Geográficas',
    '9. Demoras y Análisis de Tiempos',
    '10. Puntos de Interés',
    '11. Indicadores y Dashboard',
    '12. Historial y Animación de Rutas',
    '13. Configuración y Preferencias',
    '14. Infraestructura y Despliegue',
    '15. Stack Tecnológico',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 1. INTRODUCCIÓN
# ══════════════════════════════════════════════════════════════
doc.add_heading('1. Introducción', level=1)

doc.add_paragraph(
    'TrackMóvil es una plataforma web de rastreo vehicular en tiempo real '
    'diseñada para la gestión integral de flotas de distribución. El sistema '
    'permite monitorear la ubicación y el estado de los vehículos (móviles), '
    'gestionar pedidos y servicios, analizar demoras por zona geográfica, y '
    'proporcionar indicadores operativos clave para la toma de decisiones.'
)
doc.add_paragraph(
    'La plataforma está orientada a operaciones logísticas de distribución '
    'de gas (Riogas S.A.) e integra datos provenientes del sistema AS400 '
    'legado con tecnologías modernas de tiempo real (WebSocket) y mapas '
    'interactivos.'
)

doc.add_heading('Objetivos Principales', level=2)
objectives = [
    'Visualización en tiempo real de la flota vehicular sobre un mapa interactivo.',
    'Seguimiento del estado de pedidos y servicios vinculados a cada vehículo.',
    'Análisis de demoras y zonas de reparto para optimización logística.',
    'Generación de indicadores operativos (KPIs) en tiempo real.',
    'Herramientas de análisis histórico con animación de rutas recorridas.',
    'Gestión de zonas geográficas con polígonos y asignación de móviles.',
    'Sistema de preferencias personalizable por usuario.',
]
for obj in objectives:
    add_bullet(doc, obj)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 2. ARQUITECTURA
# ══════════════════════════════════════════════════════════════
doc.add_heading('2. Arquitectura del Sistema', level=1)

doc.add_paragraph(
    'TrackMóvil utiliza una arquitectura de tres capas con un frontend React, '
    'un Backend for Frontend (BFF) basado en API Routes de Next.js, y múltiples '
    'fuentes de datos backend.'
)

doc.add_heading('Diagrama de Componentes', level=2)
doc.add_paragraph(
    'El flujo de datos sigue el siguiente esquema:'
)
arch_items = [
    ('Frontend (React + Leaflet):', 'Renderiza mapas, tablas y modales. Se comunica '
     'exclusivamente con el BFF via API REST y recibe actualizaciones en tiempo real '
     'a través de WebSocket (Supabase Realtime).'),
    ('BFF (Next.js API Routes):', '26 grupos de endpoints que actúan como proxy '
     'entre el frontend y las APIs externas. Aplica autenticación, rate limiting, '
     'validación y transformación de datos.'),
    ('API Python (FastAPI):', 'Puente entre el sistema y la base de datos IBM AS400/DB2 '
     'vía JDBC. Expone endpoints REST para posiciones GPS, pedidos, servicios y datos '
     'de zona desde el sistema legado.'),
    ('Supabase (PostgreSQL + Realtime):', 'Base de datos moderna para coordenadas GPS, '
     'datos extendidos de móviles, puntos de interés y preferencias de usuario. '
     'Proporciona notificaciones en tiempo real vía WebSocket.'),
    ('GeneXus API:', 'Sistema de autenticación y gestión de usuarios existente.'),
]
for title, desc in arch_items:
    add_bullet(doc, desc, bold_prefix=title)

doc.add_heading('Flujo de Datos en Tiempo Real', level=2)
doc.add_paragraph(
    'Las actualizaciones GPS siguen un flujo bidireccional: los dispositivos '
    'móviles envían coordenadas → la API Python las almacena en AS400 → un '
    'proceso de sincronización las replica a Supabase → el WebSocket de Supabase '
    'notifica al frontend → los marcadores del mapa se actualizan instantáneamente. '
    'Este flujo permite latencia sub-segundo entre el movimiento real del vehículo '
    'y su representación en pantalla.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 3. AUTENTICACIÓN
# ══════════════════════════════════════════════════════════════
doc.add_heading('3. Autenticación y Control de Acceso', level=1)

doc.add_heading('Login', level=2)
doc.add_paragraph(
    'El acceso al sistema requiere autenticación mediante credenciales de usuario '
    'verificadas contra la API de GeneXus. El proceso de login es:'
)
login_steps = [
    'El usuario ingresa sus credenciales en la pantalla de login.',
    'Las credenciales se envían al BFF que las reenvía a la API GeneXus.',
    'GeneXus valida y retorna un token de sesión junto con datos del usuario '
    '(nombre, rol, empresas permitidas, flag isRoot).',
    'La sesión se sincroniza con Supabase para habilitar suscripciones en tiempo real.',
    'El usuario es redirigido al dashboard principal.',
]
for i, step in enumerate(login_steps, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'{i}. ')
    run.bold = True
    p.add_run(step)

doc.add_heading('Control de Acceso por Empresas', level=2)
doc.add_paragraph(
    'El sistema implementa un modelo de acceso basado en empresas fleteras. '
    'Los usuarios con rol "root" (isRoot = S) tienen acceso completo a todos '
    'los datos. Los usuarios regulares solo pueden visualizar información '
    'correspondiente a las empresas fleteras que tienen asignadas en el sistema. '
    'Esta restricción se aplica a: móviles, pedidos, servicios y zonas.'
)

doc.add_heading('Rutas Protegidas', level=2)
doc.add_paragraph(
    'Toda la aplicación está envuelta en un componente de ruta protegida que '
    'verifica la existencia de una sesión válida antes de permitir el acceso '
    'al dashboard. Si no existe sesión activa, el usuario es redirigido '
    'automáticamente a la pantalla de login.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 4. MAPA INTERACTIVO
# ══════════════════════════════════════════════════════════════
doc.add_heading('4. Mapa Interactivo', level=1)

doc.add_paragraph(
    'El componente central de TrackMóvil es un mapa interactivo basado en la '
    'biblioteca Leaflet, que permite la visualización geográfica de todos los '
    'elementos operativos: vehículos, pedidos, servicios, zonas y puntos de interés.'
)

doc.add_heading('4.1 Capas Base', level=2)
doc.add_paragraph('El usuario puede seleccionar entre 6 capas base de mapa:')
add_styled_table(doc,
    ['Capa', 'Proveedor', 'Descripción'],
    [
        ['Calles', 'OpenStreetMap', 'Vista estándar de calles y carreteras'],
        ['Satélite', 'Esri World Imagery', 'Imágenes satelitales de alta resolución'],
        ['Terreno', 'OpenStreetMap', 'Vista topográfica con relieve'],
        ['CartoDB Voyager', 'CartoDB', 'Diseño minimalista y moderno'],
        ['Modo Oscuro', 'CartoDB Dark Matter', 'Tema oscuro para uso nocturno'],
        ['Modo Claro', 'CartoDB Positron', 'Tema ultra-minimalista en tonos claros'],
    ]
)

doc.add_heading('4.2 Marcadores de Vehículos', level=2)
doc.add_paragraph(
    'Cada vehículo (móvil) se representa con un marcador personalizado que '
    'comunica información visual inmediata:'
)
add_styled_table(doc,
    ['Color', 'Significado'],
    [
        ['Verde', 'Vehículo activo, ocupación < 67%'],
        ['Amarillo', 'Vehículo activo, ocupación 67-99%'],
        ['Negro', 'Vehículo activo, ocupación ≥ 100%'],
        ['Gris', 'No Activo (estado 3)'],
        ['Violeta', 'Baja Momentánea (estado 4)'],
    ]
)
doc.add_paragraph()
doc.add_paragraph(
    'Cada marcador incluye un badge numérico que indica la cantidad de pedidos/servicios '
    'asignados. Las formas son personalizables por usuario: círculo, cuadrado, triángulo, '
    'rombo, hexágono o estrella. Los marcadores también están disponibles en tres tamaños: '
    'normal, compacto y mini.'
)

doc.add_heading('4.3 Marcadores de Pedidos', level=2)
doc.add_paragraph(
    'Los pedidos pendientes se muestran como marcadores coloreados según su nivel de atraso:'
)
add_styled_table(doc,
    ['Color', 'Estado'],
    [
        ['Verde', 'En hora — tiempo restante suficiente'],
        ['Amarillo', 'Hora límite cercana — próximo a vencer'],
        ['Naranja', 'Atrasado — pasada la hora comprometida'],
        ['Rojo', 'Muy atrasado — demora significativa'],
        ['Gris', 'Sin hora definida'],
    ]
)
doc.add_paragraph()
doc.add_paragraph(
    'Los pedidos finalizados muestran un marcador verde con ✓ si fueron entregados '
    '(sub_estado 3) o rojo con ✗ si no fueron entregados.'
)

doc.add_heading('4.4 Popups Informativos', level=2)
doc.add_paragraph(
    'Al hacer clic en cualquier elemento del mapa se despliega un popup con '
    'información detallada:'
)
add_bullet(doc, 'Nombre, matrícula, estado, porcentaje de ocupación, distancia '
    'recorrida, chofer actual, terminal, historial de choferes, pedidos y servicios pendientes.',
    bold_prefix='Popup de Móvil:')
add_bullet(doc, 'ID, cliente, dirección, zona, hora comprometida, tiempo de demora, '
    'producto, importe, estado.',
    bold_prefix='Popup de Pedido:')
add_bullet(doc, 'ID, tipo de servicio, cliente, dirección, estado, hora programada.',
    bold_prefix='Popup de Servicio:')

doc.add_heading('4.5 Clustering de Marcadores', level=2)
doc.add_paragraph(
    'Cuando la densidad de marcadores es alta, el sistema agrupa automáticamente '
    'marcadores cercanos en clusters que muestran un número con la cantidad de '
    'elementos contenidos. Al hacer zoom, los clusters se desagrupan progresivamente. '
    'Esta funcionalidad es activable/desactivable desde las preferencias del usuario.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 5. FLOTA VEHICULAR
# ══════════════════════════════════════════════════════════════
doc.add_heading('5. Gestión de Flota Vehicular', level=1)

doc.add_heading('5.1 Panel Lateral (Sidebar)', level=2)
doc.add_paragraph(
    'El panel lateral izquierdo presenta una estructura jerárquica en árbol con '
    'categorías expansibles: Empresas, Móviles, Pedidos, Servicios y Puntos de '
    'Interés. Cada categoría puede ocultarse/mostrarse independientemente.'
)
doc.add_paragraph(
    'Cada tarjeta de móvil en el sidebar muestra: identificador, matrícula, estado '
    'operativo, barra visual de porcentaje de ocupación (verde/amarillo/rojo), y '
    'timestamp de la última posición GPS recibida.'
)

doc.add_heading('5.2 Posiciones GPS en Tiempo Real', level=2)
doc.add_paragraph(
    'Al iniciar la aplicación se realiza una carga inicial de las últimas posiciones '
    'conocidas de todos los móviles consultando el AS400. Posteriormente, las '
    'posiciones se actualizan en tiempo real a través de WebSocket de Supabase. '
    'Un indicador visual en la esquina del mapa muestra el estado de la conexión: '
    '"📡 Tiempo Real Activo" (verde), "Conectando..." (amarillo) o '
    '"Modo Estático" (gris).'
)

doc.add_heading('5.3 Detección de Inactividad', level=2)
doc.add_paragraph(
    'El sistema marca automáticamente como inactivos aquellos móviles cuya última '
    'señal GPS exceda un umbral configurable (por defecto 30 minutos). Los '
    'marcadores de móviles inactivos cambian su apariencia visual (gris, sin '
    'animación pulse) para diferenciarse de los activos.'
)

doc.add_heading('5.4 Filtros Avanzados', level=2)
doc.add_paragraph('Los móviles pueden filtrarse por múltiples criterios simultáneos:')
add_bullet(doc, 'Activo / No Activo / Baja Momentánea', bold_prefix='Estado:')
add_bullet(doc, 'Todos / Rangos específicos de capacidad', bold_prefix='Capacidad:')
add_bullet(doc, 'Búsqueda libre por número, nombre o matrícula', bold_prefix='Texto:')
add_bullet(doc, 'Selección masiva ("Seleccionar todos" / "Limpiar")', bold_prefix='Selección:')

doc.add_heading('5.5 Información de Sesión', level=2)
doc.add_paragraph(
    'Para cada móvil se puede consultar la sesión activa: chofer actual (nombre y '
    'teléfono), terminal Android utilizado, e historial de sesiones previas del día.'
)

doc.add_heading('5.6 Selector de Empresa Fletera', level=2)
doc.add_paragraph(
    'Un dropdown en la barra superior permite filtrar toda la información del '
    'dashboard por una o varias empresas fleteras subcontratistas. Solo visible '
    'para usuarios con acceso a más de una empresa.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 6. GESTIÓN DE PEDIDOS
# ══════════════════════════════════════════════════════════════
doc.add_heading('6. Gestión de Pedidos', level=1)

doc.add_heading('6.1 Visualización en Tiempo Real', level=2)
doc.add_paragraph(
    'Los pedidos se cargan inicialmente desde la API y se mantienen actualizados '
    'en tiempo real mediante suscripciones WebSocket a la tabla de pedidos en '
    'Supabase. Cada inserción o actualización de pedido se refleja instantáneamente '
    'tanto en el mapa como en las tablas y los indicadores del dashboard.'
)

doc.add_heading('6.2 Tabla Extendida de Pedidos', level=2)
doc.add_paragraph(
    'Un modal de pantalla completa presenta una tabla interactiva con las siguientes '
    'capacidades:'
)
add_bullet(doc, 'Pendientes / Finalizados / Sin Asignar', bold_prefix='Vistas:')
add_bullet(doc, 'Texto libre, nivel de atraso, zona, móvil, producto, pedidos sin coordenadas',
    bold_prefix='Filtros:')
add_bullet(doc, 'Delay, ID, móvil, zona, cliente, producto, importe (ascendente/descendente)',
    bold_prefix='Ordenamiento:')
add_bullet(doc, '50 registros por página', bold_prefix='Paginación:')
add_bullet(doc, 'Coloreo automático de filas según nivel de atraso',
    bold_prefix='Visualización:')

doc.add_heading('6.3 Cálculo de Atrasos', level=2)
doc.add_paragraph(
    'El sistema calcula automáticamente el nivel de demora de cada pedido '
    'comparando la hora actual contra la hora máxima de entrega comprometida. '
    'Las categorías resultantes son: En Hora (verde), Hora Límite Cercana '
    '(amarillo), Atrasado (naranja), Muy Atrasado (rojo), y Sin Hora Definida (gris).'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 7. GESTIÓN DE SERVICIOS
# ══════════════════════════════════════════════════════════════
doc.add_heading('7. Gestión de Servicios', level=1)

doc.add_paragraph(
    'Análogamente a los pedidos, el sistema gestiona servicios técnicos '
    '(instalación, revisión, mantenimiento, etc.) con las mismas capacidades '
    'de visualización en tiempo real, tabla extendida con filtros, y cálculo '
    'de atrasos. Los servicios se muestran como marcadores diferenciados en '
    'el mapa y se clasifican por tipo: URGENTE, SERVICE, NOCTURNO, entre otros.'
)
doc.add_paragraph(
    'El filtrado por tipo de servicio se aplica transversalmente a todas las '
    'funcionalidades del sistema: mapa, tablas, indicadores y zonas.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 8. SISTEMA DE ZONAS
# ══════════════════════════════════════════════════════════════
doc.add_heading('8. Sistema de Zonas Geográficas', level=1)

doc.add_paragraph(
    'TrackMóvil implementa un sistema completo de zonas geográficas representadas '
    'como polígonos sobre el mapa. Las zonas soportan formato GeoJSON estándar '
    '(Feature, Polygon, MultiPolygon) y se gestionan desde el backend.'
)

doc.add_heading('8.1 Modos de Visualización', level=2)
doc.add_paragraph(
    'El sistema ofrece 5 modos de visualización seleccionables desde un control '
    'en el mapa:'
)
add_styled_table(doc,
    ['Modo', 'Descripción'],
    [
        ['Sin Zona', 'Vista estándar sin polígonos de zona'],
        ['Distribución', 'Polígonos con los colores de la tabla de distribución e identificador numérico'],
        ['Demoras', 'Polígonos coloreados con gradiente verde→rojo según minutos de demora'],
        ['Móviles en Zonas', 'Polígonos coloreados por cantidad de móviles en prioridad (0=rojo, 4+=verde)'],
        ['Zonas Activas', 'Polígonos verdes (zona activa) o rojos (zona inactiva)'],
    ]
)

doc.add_heading('8.2 Modal de Móviles por Zona', level=2)
doc.add_paragraph(
    'Al hacer clic en una zona en el modo "Móviles en Zonas", se abre un modal '
    'que detalla todos los móviles asignados a esa zona. El modal permite filtrar '
    'por tipo de servicio (URGENTE, SERVICE, NOCTURNO) e indica el estado de cada '
    'móvil: prioridad o tránsito, activo o inactivo.'
)

doc.add_heading('8.3 Estadísticas por Zona', level=2)
doc.add_paragraph(
    'Un modal de estadísticas presenta una tabla agregada por zona con las '
    'siguientes métricas:'
)
add_bullet(doc, 'Pedidos sin asignar por zona')
add_bullet(doc, 'Pedidos pendientes y atrasados')
add_bullet(doc, 'Porcentaje de atrasos')
add_bullet(doc, 'Pedidos entregados y no entregados')
add_bullet(doc, 'Porcentaje de cumplimiento')
add_bullet(doc, 'Demora en minutos')
add_bullet(doc, 'Cantidad de móviles en prioridad')
doc.add_paragraph(
    'La tabla permite ordenamiento por cualquier columna y filtrado por tipo '
    'de servicio.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 9. DEMORAS
# ══════════════════════════════════════════════════════════════
doc.add_heading('9. Demoras y Análisis de Tiempos', level=1)

doc.add_paragraph(
    'La capa de demoras visualiza el tiempo de atraso acumulado por zona '
    'geográfica. Cada polígono se colorea con un gradiente de verde (0 minutos) '
    'a rojo (alta demora). Opcionalmente, se muestran etiquetas con el número '
    'de zona y los minutos de demora sobre el centroide de cada polígono.'
)
doc.add_paragraph(
    'Los datos de demora se obtienen de una tabla dedicada y se actualizan '
    'periódicamente mediante polling configurable (por defecto cada 30 segundos). '
    'La opacidad de los polígonos es ajustable desde las preferencias del usuario '
    '(0% a 100%).'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 10. PUNTOS DE INTERÉS
# ══════════════════════════════════════════════════════════════
doc.add_heading('10. Puntos de Interés', level=1)

doc.add_heading('10.1 Marcadores Personalizados', level=2)
doc.add_paragraph(
    'Los usuarios pueden crear puntos de interés (POIs) personalizados '
    'directamente sobre el mapa. Cada POI incluye un nombre, un emoji '
    'representativo y una categoría. Los POIs se almacenan en la base de '
    'datos y son visibles para todos los usuarios de la misma empresa.'
)

doc.add_heading('10.2 Importación desde OpenStreetMap', level=2)
doc.add_paragraph(
    'El sistema permite la importación masiva de puntos de interés desde '
    'OpenStreetMap a través de la API Overpass. Las categorías disponibles '
    'incluyen: estaciones Riogas, instituciones gubernamentales, hospitales, '
    'bancos, entre otros. El proceso de importación detecta duplicados y '
    'reporta estadísticas de importación.'
)

doc.add_heading('10.3 Filtrado por Categoría', level=2)
doc.add_paragraph(
    'Los POIs pueden mostrarse u ocultarse por categoría individual o de '
    'forma global. La configuración de visibilidad se persiste en las '
    'preferencias del usuario.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 11. INDICADORES Y DASHBOARD
# ══════════════════════════════════════════════════════════════
doc.add_heading('11. Indicadores y Dashboard', level=1)

doc.add_heading('11.1 Indicadores KPI', level=2)
doc.add_paragraph(
    'La barra superior del dashboard muestra tres indicadores clave de '
    'rendimiento (KPIs) actualizados en tiempo real:'
)
add_styled_table(doc,
    ['Indicador', 'Descripción', 'Código de Color'],
    [
        ['Ped. sin Asig.', 'Cantidad de pedidos pendientes sin móvil asignado', 'Naranja si > 0'],
        ['Entregados', 'Cantidad de pedidos finalizados con entrega confirmada', 'Verde'],
        ['% Entregados', 'Porcentaje de pedidos entregados sobre total finalizados', 'Verde (>80%), Naranja (50-80%), Rojo (<50%)'],
    ]
)

doc.add_heading('11.2 Ranking / Leaderboard', level=2)
doc.add_paragraph(
    'Un modal de ranking presenta una tabla competitiva con métricas por '
    'vehículo: pedidos atrasados, pendientes, entregados, no entregados, '
    'porcentaje de cumplimiento y porcentaje de cumplimiento en hora. '
    'Los tres primeros vehículos reciben medallas (🥇 🥈 🥉). Al hacer clic '
    'en cualquier estadística, se abre la tabla de pedidos pre-filtrada por '
    'ese móvil. El ranking está disponible tanto para pedidos como para servicios.'
)

doc.add_heading('11.3 Indicador de Conexión', level=2)
doc.add_paragraph(
    'Un badge flotante en el mapa indica el estado de la conexión en tiempo '
    'real: activa (verde), reconectando (amarillo) o desconectada/estática (gris). '
    'El sistema implementa reconexión automática con backoff exponencial.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 12. HISTORIAL Y ANIMACIÓN
# ══════════════════════════════════════════════════════════════
doc.add_heading('12. Historial y Animación de Rutas', level=1)

doc.add_heading('12.1 Consulta de Recorrido Histórico', level=2)
doc.add_paragraph(
    'El sistema permite consultar y visualizar el recorrido completo de '
    'cualquier vehículo en una fecha específica. El usuario selecciona un '
    'móvil (con búsqueda por ID, nombre o matrícula) y una fecha, y el '
    'sistema descarga las coordenadas GPS históricas y dibuja la ruta '
    'completa sobre el mapa.'
)

doc.add_heading('12.2 Panel de Control de Animación', level=2)
doc.add_paragraph(
    'Un panel de control completo permite reproducir la ruta animada con '
    'las siguientes funcionalidades:'
)
add_bullet(doc, 'Play / Pause / Reset')
add_bullet(doc, '7 velocidades de reproducción: 0.1x, 0.25x, 0.5x, 1x, 2x, 5x, 10x')
add_bullet(doc, 'Filtrado por rango horario (hora inicio / hora fin)')
add_bullet(doc, 'Simplificación de ruta para mejorar legibilidad')
add_bullet(doc, 'Barra de progreso con hora actual de animación')

doc.add_heading('12.3 Comparación Dual de Rutas', level=2)
doc.add_paragraph(
    'Es posible visualizar simultáneamente las rutas de dos vehículos en '
    'la misma fecha, con colores diferenciados (azul y naranja) y un '
    'timeline unificado. Esta funcionalidad facilita la comparación de '
    'recorridos y la detección de ineficiencias en rutas paralelas.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 13. CONFIGURACIÓN
# ══════════════════════════════════════════════════════════════
doc.add_heading('13. Configuración y Preferencias', level=1)

doc.add_paragraph(
    'El sistema ofrece un amplio conjunto de preferencias personalizables '
    'que se persisten tanto localmente como en el servidor para sincronización '
    'entre dispositivos.'
)

add_styled_table(doc,
    ['Categoría', 'Preferencia', 'Opciones'],
    [
        ['Mapa', 'Capa base', 'Calles / Satélite / Terreno / CartoDB / Oscuro / Claro'],
        ['Mapa', 'Opacidad de zonas', '0% a 100%'],
        ['Mapa', 'Etiquetas de demora', 'Activar / Desactivar'],
        ['Móviles', 'Solo activos', 'Activar / Desactivar'],
        ['Móviles', 'Retraso máx. GPS', 'Minutos (predeterminado: 30)'],
        ['Marcadores', 'Tamaño de marcador', 'Normal / Compacto / Mini'],
        ['Marcadores', 'Forma', 'Círculo / Cuadrado / Triángulo / Rombo / Hexágono / Estrella'],
        ['Pedidos', 'Clustering', 'Activar / Desactivar'],
        ['Tiempo Real', 'WebSocket', 'Activar / Desactivar'],
        ['Tiempo Real', 'Polling demoras', 'Segundos (predeterminado: 30)'],
        ['Tiempo Real', 'Polling móviles/zonas', 'Segundos (predeterminado: 30)'],
        ['Visibilidad', 'Capas visibles', 'Móviles / Pedidos / Servicios / POIs'],
    ]
)

doc.add_paragraph()
doc.add_paragraph(
    'Adicionalmente, el sistema incluye un tour interactivo guiado para '
    'nuevos usuarios y una guía visual de colores e íconos accesible desde '
    'la interfaz.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 14. INFRAESTRUCTURA
# ══════════════════════════════════════════════════════════════
doc.add_heading('14. Infraestructura y Despliegue', level=1)

doc.add_heading('14.1 Contenedorización (Docker)', level=2)
doc.add_paragraph(
    'La aplicación se empaqueta como imagen Docker multi-stage optimizada '
    '(basada en node:20-alpine). El build se divide en tres etapas: instalación '
    'de dependencias, compilación de Next.js, y runtime de producción con '
    'usuario no-root por seguridad.'
)

doc.add_heading('14.2 Orquestación (Docker Compose)', level=2)
doc.add_paragraph(
    'Docker Compose gestiona dos servicios: la aplicación Next.js (puerto 3000) '
    'y la API Python AS400 (puerto 8000, activable mediante perfil). Incluye '
    'healthcheck automático cada 30 segundos.'
)

doc.add_heading('14.3 Gestión de Procesos (PM2)', level=2)
doc.add_paragraph(
    'En producción sin Docker, la aplicación se ejecuta bajo PM2 con '
    'reinicio automático, backoff exponencial, límite de memoria de 2 GB, '
    'y gestión centralizada de logs.'
)

doc.add_heading('14.4 Reverse Proxy (Nginx)', level=2)
doc.add_paragraph(
    'Nginx se utiliza como reverse proxy frente a la aplicación Next.js, '
    'manejando terminación SSL, compresión y balanceo de carga.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 15. STACK TECNOLÓGICO
# ══════════════════════════════════════════════════════════════
doc.add_heading('15. Stack Tecnológico', level=1)

add_styled_table(doc,
    ['Capa', 'Tecnología', 'Versión'],
    [
        ['Framework', 'Next.js (App Router)', '16.1.6'],
        ['UI', 'React + TypeScript', '19.1.0 / 5.x'],
        ['Estilos', 'Tailwind CSS', '4.x'],
        ['Mapas', 'Leaflet + React-Leaflet', '1.9.4 / 5.0.0'],
        ['Clustering', 'leaflet.markercluster', '1.5.3'],
        ['Animaciones', 'Framer Motion', '12.x'],
        ['Base de Datos', 'Supabase (PostgreSQL)', 'Cloud'],
        ['Tiempo Real', 'Supabase Realtime (WebSocket)', 'Cloud'],
        ['Autenticación', 'GeneXus API', 'Legacy'],
        ['Virtualización', 'react-window', '2.x'],
        ['Tour Interactivo', 'driver.js', '—'],
        ['Legacy Backend', 'FastAPI + JayDeBeAPI (Python → AS400/DB2)', '—'],
        ['Contenedores', 'Docker + Docker Compose', '—'],
        ['Procesos', 'PM2', '—'],
        ['Proxy', 'Nginx', '—'],
        ['Package Manager', 'pnpm', '—'],
    ]
)

doc.add_paragraph()
doc.add_paragraph()

# ── Pie de documento ──────────────────────────────────────────
doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('— Fin del Documento —')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.italic = True

# ── Guardar ───────────────────────────────────────────────────
output_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'docs', 'TrackMovil_Documentacion.docx')
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print(f'✅ Documento generado: {output_path}')
