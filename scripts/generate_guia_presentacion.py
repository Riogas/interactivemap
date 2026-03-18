#!/usr/bin/env python3
"""
Genera una Guía de Presentación Ejecutiva de TrackMóvil en formato .docx
Orientada a directivos — funcionalidad por funcionalidad, de mayor a menor importancia.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from datetime import datetime

doc = Document()

# ══════════════════════════════════════════════════════════════
# ESTILOS GLOBALES
# ══════════════════════════════════════════════════════════════
BLUE = RGBColor(0x1a, 0x56, 0xdb)
DARK = RGBColor(0x2d, 0x2d, 0x2d)
GRAY = RGBColor(0x66, 0x66, 0x66)
LIGHT_GRAY = RGBColor(0x99, 0x99, 0x99)
GREEN = RGBColor(0x16, 0xa3, 0x4a)
ORANGE = RGBColor(0xea, 0x58, 0x0c)
RED = RGBColor(0xdc, 0x26, 0x26)
WHITE = RGBColor(0xff, 0xff, 0xff)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = DARK

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.color.rgb = BLUE
    h.font.name = 'Calibri'
    if level == 1:
        h.font.size = Pt(22)
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(8)
    elif level == 2:
        h.font.size = Pt(16)
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(6)
    elif level == 3:
        h.font.size = Pt(13)
        h.paragraph_format.space_before = Pt(10)
        h.paragraph_format.space_after = Pt(4)


# ── Funciones auxiliares ──────────────────────────────────────
def add_styled_table(doc, headers, rows):
    """Tabla profesional con estilo Light Grid Accent 1."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    return table


def add_bullet(doc, text, bold_prefix=None):
    """Viñeta con prefijo en negrita opcional."""
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(f' {text}')
    else:
        p.add_run(text)
    return p


def add_numbered(doc, text, number):
    """Párrafo con número en negrita."""
    p = doc.add_paragraph()
    run = p.add_run(f'{number}. ')
    run.bold = True
    run.font.color.rgb = BLUE
    p.add_run(text)
    return p


def add_highlight_box(doc, text, emoji='💡'):
    """Párrafo destacado (simulación de cuadro informativo)."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f'{emoji} ')
    run.font.size = Pt(11)
    run2 = p.add_run(text)
    run2.italic = True
    run2.font.color.rgb = GRAY
    return p


def add_section_intro(doc, text):
    """Párrafo introductorio de sección con mayor tamaño."""
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(8)
    for run in p.runs:
        run.font.size = Pt(11)
    return p


def add_valor_negocio(doc, items):
    """Cuadro de valor de negocio con ítems."""
    p = doc.add_paragraph()
    run = p.add_run('Valor para el negocio:')
    run.bold = True
    run.font.color.rgb = GREEN
    run.font.size = Pt(11)
    for item in items:
        add_bullet(doc, item)


# ══════════════════════════════════════════════════════════════
# PORTADA
# ══════════════════════════════════════════════════════════════
for _ in range(5):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('TrackMóvil')
run.font.size = Pt(48)
run.font.color.rgb = BLUE
run.bold = True

doc.add_paragraph()

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Guía de Presentación Ejecutiva')
run.font.size = Pt(22)
run.font.color.rgb = GRAY

doc.add_paragraph()

desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = desc.add_run('Sistema de Rastreo Vehicular en Tiempo Real\npara la Gestión de Flotas de Distribución')
run.font.size = Pt(14)
run.font.color.rgb = LIGHT_GRAY

for _ in range(3):
    doc.add_paragraph()

meta_lines = [
    ('Empresa:', 'Riogas S.A.'),
    ('Versión:', '1.0'),
    ('Fecha:', datetime.now().strftime('%B %Y').capitalize()),
    ('Destinatarios:', 'Directorio / Gerencia'),
    ('Clasificación:', 'Uso Interno'),
]
for label, value in meta_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(label + ' ')
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = BLUE
    run2 = p.add_run(value)
    run2.font.size = Pt(12)
    run2.font.color.rgb = DARK

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# ÍNDICE
# ══════════════════════════════════════════════════════════════
doc.add_heading('Índice de Contenidos', level=1)

toc_items = [
    ('1.', 'Resumen Ejecutivo'),
    ('2.', 'Monitoreo en Tiempo Real de la Flota', '★★★'),
    ('3.', 'Gestión de Pedidos y Entregas', '★★★'),
    ('4.', 'Gestión de Servicios Técnicos', '★★★'),
    ('5.', 'Indicadores Operativos (KPIs)', '★★★'),
    ('6.', 'Sistema de Zonas Geográficas', '★★☆'),
    ('7.', 'Análisis de Demoras por Zona', '★★☆'),
    ('8.', 'Ranking de Rendimiento (Leaderboard)', '★★☆'),
    ('9.', 'Historial y Animación de Recorridos', '★★☆'),
    ('10.', 'Vista de Zonas por Móvil', '★☆☆'),
    ('11.', 'Puntos de Interés Personalizados', '★☆☆'),
    ('12.', 'Personalización y Preferencias', '★☆☆'),
    ('13.', 'Seguridad y Control de Acceso', '★☆☆'),
    ('14.', 'Arquitectura Técnica y Despliegue', '★☆☆'),
]

for item in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f'{item[0]} ')
    run.bold = True
    run.font.color.rgb = BLUE
    run.font.size = Pt(12)
    run2 = p.add_run(item[1])
    run2.font.size = Pt(12)
    if len(item) > 2:
        run3 = p.add_run(f'  {item[2]}')
        run3.font.size = Pt(10)
        run3.font.color.rgb = LIGHT_GRAY

doc.add_paragraph()
add_highlight_box(doc, '★★★ = Funcionalidad crítica   ★★☆ = Funcionalidad importante   ★☆☆ = Funcionalidad complementaria', '📊')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 1. RESUMEN EJECUTIVO
# ══════════════════════════════════════════════════════════════
doc.add_heading('1. Resumen Ejecutivo', level=1)

doc.add_paragraph(
    'TrackMóvil es una plataforma web de rastreo vehicular en tiempo real '
    'desarrollada específicamente para la operación logística de distribución '
    'de gas de Riogas S.A. El sistema permite visualizar, controlar y optimizar '
    'toda la cadena de distribución desde una única interfaz.'
)

doc.add_heading('¿Qué problema resuelve?', level=2)
doc.add_paragraph(
    'Antes de TrackMóvil, la operación de distribución dependía de sistemas legacy '
    '(AS400 + GeneXus) sin interfaz visual unificada. Los supervisores no podían '
    'ver en tiempo real dónde estaba cada vehículo, cuántos pedidos tenía asignados, '
    'ni si las entregas se estaban haciendo en hora. TrackMóvil elimina esa '
    'brecha al ofrecer una visión completa, en vivo, de toda la operación.'
)

doc.add_heading('Capacidades principales', level=2)
capabilities = [
    ('Mapa en vivo:', 'Todos los vehículos visibles en un mapa interactivo, actualizándose segundo a segundo.'),
    ('Pedidos y Servicios:', 'Seguimiento completo del estado de cada entrega y servicio técnico.'),
    ('KPIs automáticos:', 'Indicadores de cumplimiento, atrasos y rendimiento calculados al instante.'),
    ('Zonas de reparto:', 'Visualización geográfica de zonas con estadísticas de demora y ocupación.'),
    ('Ranking competitivo:', 'Tabla de rendimiento por vehículo para incentivar eficiencia.'),
    ('Historial de rutas:', 'Reproducción animada del recorrido de cualquier vehículo en cualquier fecha.'),
    ('Multi-empresa:', 'Acceso segmentado por empresa fletera con control de permisos por usuario.'),
]
for title, desc in capabilities:
    add_bullet(doc, desc, bold_prefix=title)

doc.add_heading('Datos clave', level=2)
add_styled_table(doc,
    ['Métrica', 'Valor'],
    [
        ['Actualización de posición', 'Tiempo real (< 1 segundo)'],
        ['Empresas fleteras soportadas', 'Todas las operativas en Riogas'],
        ['Tipos de entrega', 'Pedidos de gas + Servicios técnicos'],
        ['Plataforma', 'Web (accesible desde cualquier navegador)'],
        ['Disponibilidad', '24/7 vía servidor dedicado'],
        ['Integración', 'AS400 (ERP) + Supabase (tiempo real) + GeneXus (autenticación)'],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 2. MONITOREO EN TIEMPO REAL  ★★★
# ══════════════════════════════════════════════════════════════
doc.add_heading('2. Monitoreo en Tiempo Real de la Flota', level=1)
add_highlight_box(doc, 'Importancia: CRÍTICA — Es la funcionalidad central del sistema.', '🔴')

add_section_intro(doc,
    'TrackMóvil muestra la ubicación exacta de cada vehículo de la flota sobre un '
    'mapa interactivo. Las posiciones se actualizan automáticamente en tiempo real '
    'mediante tecnología WebSocket, sin necesidad de recargar la página.'
)

doc.add_heading('Mapa interactivo', level=2)
doc.add_paragraph(
    'El mapa ocupa la mayor parte de la pantalla y ofrece seis estilos de visualización '
    '(Calles, Satélite, Terreno, CartoDB, Modo Oscuro, Modo Claro). El usuario puede '
    'hacer zoom, arrastrar, y hacer clic en cualquier elemento para obtener información '
    'detallada.'
)

doc.add_heading('Identificación visual instantánea', level=2)
doc.add_paragraph(
    'Cada vehículo en el mapa es un marcador cuyo color comunica su estado operativo '
    'sin necesidad de abrir menús:'
)
add_styled_table(doc,
    ['Color del vehículo', 'Significado', 'Qué implica'],
    [
        ['🟢 Verde', 'Disponibilidad alta (< 67% cargado)', 'Puede aceptar más pedidos'],
        ['🟡 Amarillo', 'Casi lleno (67-99% cargado)', 'Capacidad limitada'],
        ['⚫ Negro', 'Totalmente cargado (100%)', 'No puede recibir más pedidos'],
        ['⚪ Gris', 'No Activo', 'Fuera de servicio / sin señal GPS'],
        ['🟣 Violeta', 'Baja Momentánea', 'Temporalmente fuera de operación'],
    ]
)

doc.add_paragraph()
doc.add_paragraph(
    'Cada marcador muestra un badge numérico con la cantidad de pedidos asignados. '
    'Los marcadores pueden personalizarse en forma (círculo, cuadrado, triángulo, '
    'rombo, hexágono, estrella) y tamaño (normal, compacto, mini).'
)

doc.add_heading('Ficha del vehículo (Popup)', level=2)
doc.add_paragraph(
    'Al hacer clic en un vehículo, se abre una ficha completa con toda la información '
    'operativa:'
)
add_styled_table(doc,
    ['Dato', 'Descripción'],
    [
        ['Identificación', 'Número de móvil, matrícula, empresa fletera'],
        ['Estado', 'Activo / No Activo / Baja Momentánea'],
        ['Capacidad', 'Barra visual de ocupación (pedidos asignados / tamaño de lote)'],
        ['Posición', 'Última coordenada GPS con fecha y hora'],
        ['Distancia', 'Kilómetros recorridos en el día'],
        ['Chofer', 'Nombre del chofer actual en sesión'],
        ['Terminal', 'Identificador del dispositivo Android'],
        ['Historial', 'Últimos choferes que usaron el vehículo'],
        ['Pendientes', 'Cantidad de pedidos y servicios sin completar'],
        ['Zonas asignadas', 'Botón para ver todas las zonas asignadas al móvil'],
    ]
)

doc.add_paragraph()
doc.add_paragraph(
    'Desde la ficha se puede: ver el recorrido histórico del día, mostrar los '
    'pedidos/servicios pendientes en el mapa, enviar un mensaje SMS al chofer, '
    'y consultar las zonas asignadas.'
)

doc.add_heading('Panel lateral de flota', level=2)
doc.add_paragraph(
    'El costado izquierdo de la pantalla presenta un panel desplegable con la lista '
    'completa de vehículos organizada en un árbol jerárquico. Cada tarjeta de móvil '
    'muestra: estado, matrícula, barra de ocupación y última posición GPS. '
    'Incluye búsqueda rápida, filtros por estado y capacidad, y selección masiva.'
)

doc.add_heading('Indicador de conexión', level=2)
doc.add_paragraph(
    'Un badge en el mapa indica el estado de la conexión en tiempo real: '
    '"📡 Tiempo Real Activo" (verde), "Conectando..." (amarillo) o '
    '"Modo Estático" (gris). El sistema se reconecta automáticamente si se '
    'pierde la conexión.'
)

add_valor_negocio(doc, [
    'Visibilidad total de la flota en todo momento.',
    'Detección inmediata de vehículos detenidos o fuera de zona.',
    'Capacidad de reacción ante incidentes logísticos en segundos.',
    'Reducción de llamadas telefónicas a choferes para consultar ubicación.',
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 3. GESTIÓN DE PEDIDOS  ★★★
# ══════════════════════════════════════════════════════════════
doc.add_heading('3. Gestión de Pedidos y Entregas', level=1)
add_highlight_box(doc, 'Importancia: CRÍTICA — Gestión completa del ciclo de entrega de gas.', '🔴')

add_section_intro(doc,
    'Cada pedido de distribución de gas se visualiza en el mapa, en tablas y en '
    'indicadores del dashboard. El sistema recibe actualizaciones en tiempo real: '
    'cuando un chofer marca un pedido como entregado, el cambio se refleja al '
    'instante en toda la plataforma.'
)

doc.add_heading('Visualización en el mapa', level=2)
doc.add_paragraph(
    'Los pedidos pendientes aparecen como marcadores en la dirección del cliente, '
    'coloreados según su nivel de atraso:'
)
add_styled_table(doc,
    ['Color', 'Estado de entrega', 'Significado operativo'],
    [
        ['🟢 Verde', 'En hora', 'Entrega dentro del plazo comprometido'],
        ['🟡 Amarillo', 'Hora límite cercana', 'Quedan pocos minutos para vencer'],
        ['🟠 Naranja', 'Atrasado', 'Ya se superó la hora prometida al cliente'],
        ['🔴 Rojo', 'Muy atrasado', 'Demora significativa, requiere atención urgente'],
        ['⚪ Gris', 'Sin hora definida', 'No tiene hora comprometida asignada'],
    ]
)

doc.add_paragraph()
doc.add_paragraph(
    'Los pedidos finalizados se muestran con marcador verde (✓ entregado) o rojo '
    '(✗ no entregado) según el sub-estado reportado.'
)

doc.add_heading('Tabla extendida de pedidos', level=2)
doc.add_paragraph(
    'Un modal de pantalla completa presenta todos los pedidos en una tabla '
    'interactiva con las siguientes capacidades:'
)
add_bullet(doc, 'Pendientes / Finalizados / Sin Asignar / Entregados', bold_prefix='Vistas:')
add_bullet(doc, 'Texto libre, nivel de atraso, zona, móvil, producto', bold_prefix='Filtros:')
add_bullet(doc, 'Por todas las columnas (ID, móvil, zona, cliente, dirección, hora máxima, '
    'producto, importe, estado, observaciones, demora) en orden ascendente y descendente', bold_prefix='Ordenamiento:')
add_bullet(doc, '50 registros por página con navegación rápida', bold_prefix='Paginación:')
add_bullet(doc, 'Coloreo automático de filas según nivel de demora', bold_prefix='Visual:')

doc.add_heading('Popup de detalle de pedido', level=2)
doc.add_paragraph(
    'Al hacer clic en un pedido del mapa, se despliega toda la información: ID, '
    'cliente, dirección, zona, producto, cantidad, importe, hora comprometida, '
    'hora real de entrega, minutos de demora, estado y observaciones del cliente.'
)

add_valor_negocio(doc, [
    'Control en tiempo real de cada entrega de gas.',
    'Identificación inmediata de pedidos atrasados para tomar acción.',
    'Vista consolidada para supervisores: ¿cuántos pedidos faltan? ¿cuáles están en riesgo?',
    'Detección de pedidos sin asignar que necesitan un vehículo.',
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 4. GESTIÓN DE SERVICIOS  ★★★
# ══════════════════════════════════════════════════════════════
doc.add_heading('4. Gestión de Servicios Técnicos', level=1)
add_highlight_box(doc, 'Importancia: CRÍTICA — Seguimiento de servicios de instalación, revisión y mantenimiento.', '🔴')

add_section_intro(doc,
    'Además de la distribución de gas, Riogas realiza servicios técnicos en '
    'domicilio (instalaciones, revisiones, mantenimientos). TrackMóvil gestiona '
    'estos servicios con la misma profundidad que los pedidos.'
)

doc.add_paragraph(
    'Los servicios comparten la misma infraestructura de tiempo real, el mismo '
    'sistema de colores por atraso, y las mismas capacidades de tabla extendida '
    'con filtros y ordenamiento. Se clasifican por tipo: URGENTE, SERVICE, NOCTURNO, '
    'lo que permite priorizar la operación.'
)

doc.add_heading('Diferencias con pedidos', level=2)
add_styled_table(doc,
    ['Aspecto', 'Pedidos', 'Servicios'],
    [
        ['Objeto', 'Entrega de gas (producto físico)', 'Servicio técnico (trabajo en campo)'],
        ['Tipos', 'Único', 'URGENTE / SERVICE / NOCTURNO'],
        ['Marcador', 'Ícono de caja 📦', 'Ícono de herramienta 🔧'],
        ['Métricas', 'Producto, cantidad, importe', 'Tipo de servicio, descripción'],
        ['Tiempo real', '✅ Sí', '✅ Sí'],
        ['Tabla extendida', '✅ Sí, con 12 columnas ordenables', '✅ Sí, con 12 columnas ordenables'],
    ]
)

add_valor_negocio(doc, [
    'Visión unificada de pedidos Y servicios en la misma plataforma.',
    'Priorización automática de servicios urgentes.',
    'Control de atrasos en servicios técnicos igual que en entregas de gas.',
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 5. INDICADORES KPIs  ★★★
# ══════════════════════════════════════════════════════════════
doc.add_heading('5. Indicadores Operativos (KPIs)', level=1)
add_highlight_box(doc, 'Importancia: CRÍTICA — Métricas de rendimiento en tiempo real para la toma de decisiones.', '🔴')

add_section_intro(doc,
    'La barra superior del dashboard muestra indicadores clave de rendimiento (KPIs) '
    'que se actualizan automáticamente en tiempo real. Estos indicadores permiten a '
    'los supervisores evaluar el estado de la operación de un vistazo.'
)

doc.add_heading('Indicadores de pedidos', level=2)
add_styled_table(doc,
    ['Indicador', 'Qué muestra', 'Código de color'],
    [
        ['📦 Sin Asignar', 'Pedidos pendientes sin móvil asignado', 'Naranja si > 0, verde si = 0'],
        ['✅ Entregados', 'Pedidos finalizados con entrega confirmada (sub-estado 3 y 16)', 'Verde'],
        ['📊 % Entregados', 'Porcentaje de pedidos entregados sobre total finalizados', 'Verde (>80%), Naranja (50-80%), Rojo (<50%)'],
    ]
)

doc.add_heading('Indicadores de servicios', level=2)
add_styled_table(doc,
    ['Indicador', 'Qué muestra'],
    [
        ['🔧 Pendientes', 'Cantidad de servicios técnicos aún no realizados'],
        ['✅ Realizados', 'Cantidad de servicios completados hoy'],
        ['⚠️ Atrasados', 'Servicios que superaron la hora comprometida'],
    ]
)

doc.add_heading('Interactividad', level=2)
doc.add_paragraph(
    'Los indicadores son clickeables: al hacer clic en "Sin Asignar" se abre '
    'automáticamente la tabla de pedidos filtrada por pedidos sin móvil. Al hacer '
    'clic en "Entregados" se muestra la vista de entregas confirmadas. Esto permite '
    'pasar de la visión global al detalle en un solo clic.'
)

add_valor_negocio(doc, [
    'Visión instantánea del estado de la operación.',
    'Detección inmediata de problemas (pedidos sin asignar, altos porcentajes de atraso).',
    'Permite tomar decisiones rápidas basadas en datos reales, no en estimaciones.',
    'Indicadores enlazados al detalle para investigar cualquier anomalía.',
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 6. SISTEMA DE ZONAS  ★★☆
# ══════════════════════════════════════════════════════════════
doc.add_heading('6. Sistema de Zonas Geográficas', level=1)
add_highlight_box(doc, 'Importancia: ALTA — Herramienta clave para la planificación de la distribución.', '🟠')

add_section_intro(doc,
    'TrackMóvil implementa un sistema de zonas geográficas representadas como '
    'polígonos sobre el mapa. Estas zonas corresponden a las áreas de reparto '
    'definidas operativamente y permiten analizar la operación desde una '
    'perspectiva geográfica.'
)

doc.add_heading('Modos de visualización', level=2)
doc.add_paragraph(
    'El sistema ofrece cinco modos de zona seleccionables desde un control en el mapa:'
)
add_styled_table(doc,
    ['Modo', 'Qué muestra', 'Para qué sirve'],
    [
        ['Sin Zona', 'Mapa limpio sin polígonos', 'Vista estándar sin distracciones'],
        ['Distribución', 'Polígonos con colores identificatorios y número de zona', 'Identificar las zonas de reparto geográficamente'],
        ['Demoras', 'Polígonos teñidos verde→rojo según minutos de demora', 'Detectar zonas con problemas de demora'],
        ['Móviles en Zonas', 'Polígonos coloreados por cantidad de móviles en prioridad', 'Ver si las zonas tienen cobertura suficiente'],
        ['Zonas Activas', 'Polígonos verdes (activa) o rojos (inactiva) con leyenda', 'Saber cuáles zonas están operativas hoy'],
    ]
)

doc.add_heading('Estadísticas por zona', level=2)
doc.add_paragraph(
    'Un modal de estadísticas presenta una tabla agregada por zona con métricas como: '
    'pedidos sin asignar, pedidos pendientes, pedidos atrasados, porcentaje de atraso, '
    'entregas confirmadas, entregas fallidas, porcentaje de cumplimiento, demora en '
    'minutos y cantidad de móviles en prioridad. La tabla permite ordenar por '
    'cualquier columna y filtrar por tipo de servicio.'
)

add_valor_negocio(doc, [
    'Visión geográfica del rendimiento operativo.',
    'Identificación de zonas con cobertura insuficiente de vehículos.',
    'Detección de zonas con demoras recurrentes para optimizar rutas.',
    'Herramienta de planificación para redistribuir recursos.',
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 7. ANÁLISIS DE DEMORAS  ★★☆
# ══════════════════════════════════════════════════════════════
doc.add_heading('7. Análisis de Demoras por Zona', level=1)
add_highlight_box(doc, 'Importancia: ALTA — Detección visual de zonas con problemas de tiempo.', '🟠')

add_section_intro(doc,
    'La capa de demoras visualiza el tiempo de atraso acumulado por zona geográfica. '
    'Cada polígono se colorea con un gradiente de verde (0 minutos de demora) a rojo '
    '(alta demora), permitiendo identificar "zonas calientes" de un vistazo.'
)

doc.add_paragraph(
    'Opcionalmente, cada zona muestra una etiqueta flotante con su número y los '
    'minutos de demora. Los datos se actualizan periódicamente (cada 30 segundos '
    'por defecto, configurable). La opacidad de los polígonos es ajustable desde '
    'las preferencias del usuario (0% a 100%).'
)

add_valor_negocio(doc, [
    'Mapa de calor instantáneo de demoras operativas.',
    'Permite priorizar intervenciones en las zonas más problemáticas.',
    'Base objetiva para reuniones de análisis de eficiencia.',
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 8. RANKING / LEADERBOARD  ★★☆
# ══════════════════════════════════════════════════════════════
doc.add_heading('8. Ranking de Rendimiento (Leaderboard)', level=1)
add_highlight_box(doc, 'Importancia: ALTA — Herramienta de gestión del rendimiento de la flota.', '🟠')

add_section_intro(doc,
    'El Leaderboard es un ranking competitivo que ordena todos los vehículos por '
    'su rendimiento operativo del día. Accesible desde el botón 🏆 en la barra '
    'lateral del mapa.'
)

doc.add_heading('Tarjetas de resumen', level=2)
add_styled_table(doc,
    ['Tarjeta', 'Valor que muestra'],
    [
        ['Móviles', 'Total de vehículos activos'],
        ['Entregas', 'Total de pedidos + servicios entregados/realizados'],
        ['En Hora', 'Entregas completadas dentro del horario comprometido'],
        ['Cumplimiento', 'Porcentaje general de entregas a tiempo'],
    ]
)

doc.add_heading('Tabla de ranking por vehículo', level=2)
doc.add_paragraph(
    'Cada fila muestra un vehículo con: posición en el ranking (#1, #2, #3 con '
    'medallas 🥇🥈🥉), cantidad de entregas, porcentaje de cumplimiento (barra '
    'visual con colores verde/amarillo/rojo), entregas en hora, pendientes, y '
    'total asignado. La tabla es ordenable por cualquier columna.'
)
doc.add_paragraph(
    'Al hacer clic en cualquier estadística del ranking, se abre la tabla de '
    'pedidos pre-filtrada por ese vehículo, permitiendo investigar el detalle '
    'de su rendimiento.'
)

add_valor_negocio(doc, [
    'Visión comparativa del rendimiento de toda la flota.',
    'Incentivo competitivo: los choferes pueden ver quién tiene mejor cumplimiento.',
    'Identificación de móviles con bajo rendimiento para tomar acciones correctivas.',
    'Métricas objetivas de eficiencia por vehículo (no subjetivas).',
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 9. HISTORIAL Y ANIMACIÓN  ★★☆
# ══════════════════════════════════════════════════════════════
doc.add_heading('9. Historial y Animación de Recorridos', level=1)
add_highlight_box(doc, 'Importancia: ALTA — Análisis forense y verificación de rutas.', '🟠')

add_section_intro(doc,
    'TrackMóvil permite consultar y reproducir el recorrido completo de cualquier '
    'vehículo en cualquier fecha pasada. Las coordenadas GPS históricas se recuperan '
    'y se dibujan como una ruta animada sobre el mapa.'
)

doc.add_heading('Control de reproducción', level=2)
add_bullet(doc, 'Play / Pause / Reset', bold_prefix='Controles:')
add_bullet(doc, '0.1x, 0.25x, 0.5x, 1x, 2x, 5x, 10x', bold_prefix='Velocidades:')
add_bullet(doc, 'Establecer hora inicio y hora fin para enfocarse en un tramo', bold_prefix='Rango horario:')
add_bullet(doc, 'Barra de progreso con hora actual de la animación', bold_prefix='Timeline:')
add_bullet(doc, 'Reducción automática de puntos GPS para mejorar visualización', bold_prefix='Simplificación:')

doc.add_heading('Comparación dual de rutas', level=2)
doc.add_paragraph(
    'Es posible visualizar simultáneamente las rutas de dos vehículos en la misma '
    'fecha, con colores diferenciados (azul y naranja) y timeline unificado. '
    'Esto facilita comparar recorridos y detectar ineficiencias en rutas paralelas.'
)

add_valor_negocio(doc, [
    'Verificar que un vehículo realmente visitó la dirección del cliente.',
    'Reconstruir la secuencia de eventos ante un reclamo.',
    'Detectar desvíos de ruta o paradas no autorizadas.',
    'Comparar eficiencia de rutas entre dos vehículos.',
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 10. ZONAS POR MÓVIL  ★☆☆
# ══════════════════════════════════════════════════════════════
doc.add_heading('10. Vista de Zonas por Móvil', level=1)
add_highlight_box(doc, 'Importancia: COMPLEMENTARIA — Consulta detallada de asignación de zonas.', '🟡')

add_section_intro(doc,
    'Desde la ficha de cada vehículo (popup), un botón 📍 permite ver todas las '
    'zonas geográficas asignadas a ese móvil. La información se presenta en un '
    'modal con las siguientes características:'
)

add_bullet(doc, 'URGENTE / NOCTURNO / SERVICE / Todos', bold_prefix='Filtro por tipo:')
add_bullet(doc, 'La zona principal del móvil (estrella ⭐) vs. zona de tránsito (flecha 🔵)', bold_prefix='Prioridad/Tránsito:')
add_bullet(doc, 'Nombre y número de cada zona asignada', bold_prefix='Detalle:')
add_bullet(doc, 'Cantidad de zonas por tipo de servicio', bold_prefix='Resumen:')

add_valor_negocio(doc, [
    'Consulta rápida: ¿en qué zonas opera un vehículo específico?',
    'Verificar si la distribución de zonas es equitativa.',
    'Identificar si un móvil tiene zonas de emergencia (urgente) asignadas.',
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 11. PUNTOS DE INTERÉS  ★☆☆
# ══════════════════════════════════════════════════════════════
doc.add_heading('11. Puntos de Interés Personalizados', level=1)
add_highlight_box(doc, 'Importancia: COMPLEMENTARIA — Enriquecimiento del mapa con referencias útiles.', '🟡')

add_section_intro(doc,
    'Los usuarios pueden crear marcadores personalizados (Puntos de Interés o POIs) '
    'directamente sobre el mapa. Cada POI incluye un nombre, un emoji representativo '
    'y una categoría.'
)

doc.add_heading('Creación manual', level=2)
doc.add_paragraph(
    'Clic derecho en el mapa → crear nuevo punto de interés. Se asigna nombre, '
    'categoría y emoji. Los POIs se almacenan en la base de datos y son visibles '
    'para todos los usuarios de la misma empresa.'
)

doc.add_heading('Importación desde OpenStreetMap', level=2)
doc.add_paragraph(
    'El sistema permite importar masivamente puntos de interés de categorías '
    'predefinidas: estaciones Riogas, instituciones gubernamentales, hospitales, '
    'bancos, entre otros. El proceso detecta duplicados automáticamente.'
)

add_valor_negocio(doc, [
    'Marcar ubicaciones relevantes para la operación (depósitos, clientes frecuentes, etc.).',
    'Referencia visual rápida para supervisores.',
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 12. PERSONALIZACIÓN  ★☆☆
# ══════════════════════════════════════════════════════════════
doc.add_heading('12. Personalización y Preferencias', level=1)
add_highlight_box(doc, 'Importancia: COMPLEMENTARIA — Adaptación de la interfaz a cada usuario.', '🟡')

add_section_intro(doc,
    'Cada usuario puede personalizar la interfaz según sus necesidades a través '
    'del panel de preferencias accesible desde el ícono ⚙️:'
)

add_styled_table(doc,
    ['Categoría', 'Preferencia', 'Opciones disponibles'],
    [
        ['Mapa', 'Capa base', 'Calles / Satélite / Terreno / CartoDB / Oscuro / Claro'],
        ['Mapa', 'Opacidad de zonas', '0% a 100% (slider)'],
        ['Mapa', 'Etiquetas de demora', 'Activar / Desactivar'],
        ['Móviles', 'Solo activos', 'Mostrar solo móviles con señal GPS reciente'],
        ['Móviles', 'Retraso máx. GPS', 'Minutos sin señal para marcar como inactivo'],
        ['Marcadores', 'Tamaño', 'Normal / Compacto / Mini'],
        ['Marcadores', 'Forma', 'Círculo / Cuadrado / Triángulo / Rombo / Hexágono / Estrella'],
        ['Pedidos', 'Clustering', 'Agrupar marcadores cercanos automáticamente'],
        ['Tiempo Real', 'WebSocket', 'Activar / Desactivar actualizaciones en vivo'],
        ['Tiempo Real', 'Polling demoras', 'Frecuencia de actualización (segundos)'],
        ['Visibilidad', 'Capas', 'Móviles / Pedidos / Servicios / POIs (on/off individual)'],
    ]
)

doc.add_paragraph()
doc.add_paragraph(
    'Las preferencias se guardan automáticamente y se preservan entre sesiones. '
    'Adicionalmente, el sistema incluye un tour interactivo guiado para nuevos '
    'usuarios y una guía visual de colores e iconos.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 13. SEGURIDAD  ★☆☆
# ══════════════════════════════════════════════════════════════
doc.add_heading('13. Seguridad y Control de Acceso', level=1)
add_highlight_box(doc, 'Importancia: COMPLEMENTARIA — Infraestructura de seguridad siempre activa.', '🟡')

add_section_intro(doc,
    'TrackMóvil implementa múltiples capas de seguridad para proteger los datos '
    'operativos de la flota:'
)

doc.add_heading('Autenticación', level=2)
doc.add_paragraph(
    'El acceso requiere credenciales verificadas contra la API de GeneXus (sistema '
    'de autenticación existente de Riogas). Los tokens de sesión se gestionan de '
    'forma segura y tienen expiración automática.'
)

doc.add_heading('Control de acceso por empresa', level=2)
doc.add_paragraph(
    'Los usuarios regulares solo ven datos de las empresas fleteras que tienen '
    'asignadas. Los usuarios con rol "root" tienen acceso a todas las empresas. '
    'Esta restricción se aplica transversalmente a: móviles, pedidos, servicios, '
    'zonas y estadísticas.'
)

doc.add_heading('Protaciones adicionales', level=2)
add_bullet(doc, 'Todas las páginas requieren sesión activa. Acceso sin login redirige automáticamente.', bold_prefix='Rutas protegidas:')
add_bullet(doc, 'Límite de peticiones por IP para prevenir abuso.', bold_prefix='Rate limiting:')
add_bullet(doc, 'Validación estricta de todos los datos de entrada con Zod.', bold_prefix='Validación:')
add_bullet(doc, 'Comunicación cifrada HTTPS en producción.', bold_prefix='HTTPS:')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 14. ARQUITECTURA TÉCNICA  ★☆☆
# ══════════════════════════════════════════════════════════════
doc.add_heading('14. Arquitectura Técnica y Despliegue', level=1)
add_highlight_box(doc, 'Importancia: COMPLEMENTARIA — Información técnica de respaldo.', '🟡')

doc.add_heading('Arquitectura del sistema', level=2)
doc.add_paragraph(
    'TrackMóvil utiliza una arquitectura moderna de tres capas:'
)

add_styled_table(doc,
    ['Capa', 'Tecnología', 'Función'],
    [
        ['Frontend', 'React 19 + TypeScript + Leaflet', 'Interfaz visual: mapa, tablas, modales'],
        ['BFF (Servidor)', 'Next.js 16 API Routes', 'Proxy, autenticación, rate limiting, validación'],
        ['Base de datos', 'Supabase (PostgreSQL)', 'Datos de flota + WebSocket en tiempo real'],
        ['ERP Legacy', 'AS400 / DB2 (vía API Python)', 'Datos históricos de GPS, pedidos y servicios'],
        ['Autenticación', 'GeneXus API', 'Login y gestión de usuarios existente'],
    ]
)

doc.add_heading('Flujo de datos', level=2)
doc.add_paragraph(
    'El ciclo de datos es: App Móvil del chofer → envía GPS a Supabase → '
    'WebSocket notifica al frontend → marcador del mapa se actualiza al instante. '
    'Todo el flujo ocurre en menos de 1 segundo.'
)

doc.add_heading('Integraciones', level=2)
add_bullet(doc, 'PostgreSQL en la nube, con API REST y WebSocket en tiempo real.', bold_prefix='Supabase:')
add_bullet(doc, 'Servidor ERP principal de Riogas. Se accede via API Python (FastAPI + JDBC).', bold_prefix='AS400 / DB2:')
add_bullet(doc, 'Plataforma de autenticación existente.', bold_prefix='GeneXus:')
add_bullet(doc, 'Tiles de mapa de OpenStreetMap, Esri y CartoDB.', bold_prefix='Mapas:')

doc.add_heading('Stack tecnológico', level=2)
add_styled_table(doc,
    ['Componente', 'Tecnología'],
    [
        ['Framework web', 'Next.js 16.1.6 (App Router)'],
        ['Interfaz', 'React 19 + TypeScript'],
        ['Estilos', 'Tailwind CSS 4'],
        ['Mapas', 'Leaflet + React-Leaflet'],
        ['Animaciones', 'Framer Motion'],
        ['Base de datos', 'Supabase (PostgreSQL)'],
        ['API Legacy', 'Python FastAPI + JT400 (JDBC)'],
        ['Containerización', 'Docker + Docker Compose'],
        ['Gestión de procesos', 'PM2'],
        ['Proxy reverso', 'Nginx'],
    ]
)

doc.add_heading('Despliegue', level=2)
doc.add_paragraph(
    'La aplicación se despliega como contenedor Docker en un servidor Linux, con '
    'Nginx como proxy reverso para HTTPS. Alternativamente, se puede ejecutar '
    'con PM2 directamente. El proceso de despliegue está automatizado mediante '
    'scripts de shell dedicados.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CIERRE
# ══════════════════════════════════════════════════════════════
doc.add_heading('Resumen de Funcionalidades', level=1)

doc.add_paragraph(
    'A continuación se presenta un resumen consolidado de todas las funcionalidades '
    'de TrackMóvil, ordenadas por importancia:'
)

add_styled_table(doc,
    ['#', 'Funcionalidad', 'Importancia', 'Descripción breve'],
    [
        ['1', 'Monitoreo en Tiempo Real', '★★★', 'Todos los vehículos visibles en el mapa, actualizándose al instante'],
        ['2', 'Gestión de Pedidos', '★★★', 'Control completo del ciclo de entrega de gas con alertas de atraso'],
        ['3', 'Gestión de Servicios', '★★★', 'Seguimiento de servicios técnicos con la misma profundidad que pedidos'],
        ['4', 'Indicadores KPI', '★★★', 'Métricas de rendimiento en tiempo real en la barra superior'],
        ['5', 'Zonas Geográficas', '★★☆', '5 modos de visualización de zonas de reparto con estadísticas'],
        ['6', 'Análisis de Demoras', '★★☆', 'Mapa de calor por zona con gradiente verde→rojo'],
        ['7', 'Ranking de Rendimiento', '★★☆', 'Tabla competitiva de eficiencia por vehículo'],
        ['8', 'Historial de Recorridos', '★★☆', 'Reproducción animada de rutas con 7 velocidades'],
        ['9', 'Zonas por Móvil', '★☆☆', 'Consulta de zonas asignadas a un vehículo específico'],
        ['10', 'Puntos de Interés', '★☆☆', 'Marcadores personalizados en el mapa'],
        ['11', 'Personalización', '★☆☆', 'Preferencias de visualización por usuario'],
        ['12', 'Seguridad', '★☆☆', 'Autenticación, permisos por empresa, rate limiting'],
        ['13', 'Arquitectura', '★☆☆', 'Stack moderno con integración a sistemas legacy de Riogas'],
    ]
)

doc.add_paragraph()
doc.add_paragraph()

# ── Pie de documento ──────────────────────────────────────────
doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('— Fin del Documento —')
run.font.size = Pt(10)
run.font.color.rgb = LIGHT_GRAY
run.italic = True

doc.add_paragraph()
footer2 = doc.add_paragraph()
footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer2.add_run('TrackMóvil — Riogas S.A. — Documento de uso interno')
run.font.size = Pt(9)
run.font.color.rgb = LIGHT_GRAY

# ── Guardar ───────────────────────────────────────────────────
output_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'docs', 'TrackMovil_Guia_Presentacion.docx')
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print(f'✅ Guía de presentación generada: {output_path}')
print(f'   Ubicación: docs/TrackMovil_Guia_Presentacion.docx')
